"""Document chunking for The Unofficial Guide.

Final stage of the pipeline described in planning.md:
    collect_raw.py (raw_text/) -> clean.py (cleaned/) -> ingest.py (chunks.json)

Reads the already-cleaned text in cleaned/, splits each document into overlapping
chunks with the recursive splitter (1000 chars / 150 overlap from planning.md), and
attaches metadata (source, program, person, doc_type) drawn from raw_text/manifest.json.
This module also holds the loaders (used by collect_raw.py) and clean_text (used by
clean.py). Run directly to see the chunk count and a few sample chunks:

    python ingest.py
"""

import json
import re
from pathlib import Path

DOCUMENTS_DIR = Path(__file__).parent / "documents"
CLEANED_DIR = Path(__file__).parent / "cleaned"
MANIFEST_PATH = Path(__file__).parent / "raw_text" / "manifest.json"
OUTPUT_FILE = Path(__file__).parent / "chunks.json"

# --- Chunking parameters (from the Chunking Strategy section of planning.md) ---
CHUNK_SIZE = 1000      # characters, ~250 tokens, fits the all-MiniLM-L6-v2 window
CHUNK_OVERLAP = 150    # characters, 15%, preserves prerequisite chains across boundaries

# Separator-priority list for recursive splitting: try paragraph breaks first,
# then lines, then sentences, then words, then a hard character cut as last resort.
SEPARATORS = ["\n\n", "\n", ". ", " ", ""]

# Map each source (by filename stem, so .pdf and .docx versions map the same) to
# its program and a human-readable label, so chunks can be filtered and attributed
# later (mitigates the cross-program collision risk).
SOURCE_INFO = {
    "bscs-handout": ("BSCS Handout", "CS"),
    "2024-2026-UDC-Catalog-word": ("SEAS Catalog 2024-2026", "general"),
    "Syllabi_Computer_Science": ("CS Syllabi Archive", "CS"),
    "Syllabi_Cybersecurity": ("Cybersecurity Syllabi Archive", "Cybersecurity"),
    "udc_links": ("UDC Links Master File", "general"),
}

# The catalog .docx (exported from Acrobat) supersedes the .pdf, which extracts as
# garbled interleaved columns. Skip the PDF so the catalog is not ingested twice.
SKIP_FILES = {"2024-2026-UDC-Catalog.pdf"}

# doc_type per local-file stem, for metadata filtering.
FILE_DOC_TYPE = {
    "bscs-handout": "handbook",
    "2024-2026-UDC-Catalog-word": "catalog",
    "Syllabi_Computer_Science": "syllabus",
    "Syllabi_Cybersecurity": "syllabus",
    "udc_links": "links",
}

# Program-track and department pages: slug -> (label, program).
PROGRAM_PAGE_INFO = {
    "dept-main": ("Department Main Page", "general"),
    "program-bscs": ("Program Track: BSCS", "CS"),
    "program-bs-cyber": ("Program Track: BS Cybersecurity", "Cybersecurity"),
    "program-mscs": ("Program Track: MSCS", "CS"),
    "program-ms-cyber": ("Program Track: MS Cybersecurity", "Cybersecurity"),
    "program-abm-cs": ("Program Track: ABM CS", "CS"),
    "prerequisite-map": ("Official Prerequisite Map", "general"),
}


# --------------------------------------------------------------------------- #
# Loaders
# --------------------------------------------------------------------------- #
def load_pdf(path: Path) -> str:
    """Extract text from a digitally-created PDF with pdfplumber (no OCR)."""
    try:
        import pdfplumber
    except ImportError as exc:
        raise SystemExit(
            "pdfplumber is required. Install it with: pip install -r requirements.txt"
        ) from exc

    with pdfplumber.open(path) as pdf:
        pages = [page.extract_text() for page in pdf.pages]
    return "\n\n".join(text for text in pages if text)


def load_text(path: Path) -> str:
    """Read a Markdown or plain-text file directly."""
    return path.read_text(encoding="utf-8", errors="ignore")


def load_docx(path: Path) -> str:
    """Read a Word document in true reading order, including table cells.

    Preferred over PDF for the catalog: Word stores text in logical order, so
    there is no multi-column interleaving to untangle.
    """
    try:
        import docx
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required for .docx files. "
            "Install it with: pip install -r requirements.txt"
        ) from exc

    document = docx.Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def load_html(path: Path) -> str:
    """Strip tags/boilerplate from a saved web page.

    The faculty profiles and RateMyProfessors reviews listed in planning.md are
    URLs, not local files yet. Once saved into documents/ as .html this loader
    handles them; it needs beautifulsoup4 (add it to requirements.txt first).
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise SystemExit(
            "beautifulsoup4 is required for .html files. "
            "Add it to requirements.txt and pip install."
        ) from exc

    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n")


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".md": load_text,
    ".txt": load_text,
    ".html": load_html,
}


# --------------------------------------------------------------------------- #
# Cleaning
# --------------------------------------------------------------------------- #
def clean_text(text: str) -> str:
    """Remove page-number lines and layout noise, then collapse whitespace."""
    kept = []
    for line in text.splitlines():
        stripped = line.strip()
        # Drop lines that are only a page number.
        if re.fullmatch(r"\d{1,4}", stripped):
            continue
        # Drop the catalog's letter-spaced banner lines (e.g. "U N I V E R S I T Y").
        if stripped and re.fullmatch(r"(?:[A-Za-z] )+[A-Za-z]?", stripped):
            continue
        kept.append(stripped)

    text = "\n".join(kept)
    text = re.sub(r"[ \t]+", " ", text)      # collapse runs of spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)   # collapse blank-line runs
    return text.strip()


# --------------------------------------------------------------------------- #
# Recursive chunking
# --------------------------------------------------------------------------- #
def _split_recursive(text: str, separators: list[str]) -> list[str]:
    """Break text into pieces no larger than CHUNK_SIZE, preferring early separators."""
    if len(text) <= CHUNK_SIZE:
        return [text] if text.strip() else []

    # Choose the highest-priority separator that occurs in this text.
    separator = separators[-1]
    rest = separators[-1:]
    for i, candidate in enumerate(separators):
        if candidate == "" or candidate in text:
            separator = candidate
            rest = separators[i + 1:]
            break

    # Last resort: no usable separator left, cut on character count.
    if separator == "":
        return [text[i:i + CHUNK_SIZE] for i in range(0, len(text), CHUNK_SIZE)]

    parts = text.split(separator)
    pieces: list[str] = []
    for idx, part in enumerate(parts):
        # Re-attach the separator we split on so boundaries read naturally.
        segment = part + (separator if idx < len(parts) - 1 else "")
        if not segment:
            continue
        if len(segment) <= CHUNK_SIZE:
            pieces.append(segment)
        else:
            pieces.extend(_split_recursive(segment, rest))
    return pieces


def chunk_document(text: str) -> list[str]:
    """Recursively split cleaned text into overlapping chunks of ~CHUNK_SIZE."""
    pieces = _split_recursive(text, SEPARATORS)

    chunks: list[str] = []
    current = ""
    for piece in pieces:
        if current and len(current) + len(piece) > CHUNK_SIZE:
            chunks.append(current.strip())
            # Carry the tail of the previous chunk forward as overlap.
            overlap_tail = current[-CHUNK_OVERLAP:] if CHUNK_OVERLAP else ""
            current = overlap_tail + piece
        else:
            current += piece
    if current.strip():
        chunks.append(current.strip())

    return [chunk for chunk in chunks if chunk]


# --------------------------------------------------------------------------- #
# Pipeline
# --------------------------------------------------------------------------- #
def derive_metadata(entry: dict) -> dict:
    """Build chunk metadata (source, program, person, doc_type) for a manifest entry."""
    slug = entry["slug"]
    person = entry.get("person")

    if entry["type"] == "file":
        label, program = SOURCE_INFO.get(slug, (entry["source"], "general"))
        return {"source": label, "program": program, "person": None,
                "doc_type": FILE_DOC_TYPE.get(slug, "document")}

    if slug.startswith("rmp-"):
        return {"source": f"RateMyProfessors: {person}", "program": None,
                "person": person, "doc_type": "reviews"}

    if slug.startswith("faculty-"):
        return {"source": f"Faculty Profile: {person}", "program": None,
                "person": person, "doc_type": "faculty_profile"}

    label, program = PROGRAM_PAGE_INFO.get(slug, (slug, "general"))
    return {"source": label, "program": program, "person": None, "doc_type": "program_page"}


def load_and_chunk() -> list[dict]:
    """Read each cleaned document, chunk it, and return chunk records with metadata."""
    if not MANIFEST_PATH.exists():
        raise SystemExit("No manifest found. Run collect_raw.py and clean.py first.")

    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    records: list[dict] = []

    for entry in manifest:
        cleaned_path = CLEANED_DIR / f"{entry['slug']}.txt"
        if entry["chars"] == 0 or not cleaned_path.exists():
            continue  # source failed to load or was not cleaned

        text = cleaned_path.read_text(encoding="utf-8")
        metadata = derive_metadata(entry)

        for i, chunk in enumerate(chunk_document(text)):
            records.append({
                "id": f"{entry['slug']}-{i}",
                "text": chunk,
                "metadata": {**metadata, "slug": entry["slug"], "chunk_index": i},
            })

    return records


def main() -> None:
    records = load_and_chunk()

    if not records:
        raise SystemExit("No chunks produced. Run collect_raw.py and clean.py first.")

    # Per-source counts and length stats, for the review the milestone asks for.
    by_source: dict[str, int] = {}
    lengths = []
    for record in records:
        by_source[record["metadata"]["source"]] = by_source.get(record["metadata"]["source"], 0) + 1
        lengths.append(len(record["text"]))

    print(f"Total chunks: {len(records)}")
    print(f"Chunk length: min {min(lengths)}, max {max(lengths)}, avg {sum(lengths) // len(lengths)}")
    print("\nChunks per source:")
    for source, count in sorted(by_source.items()):
        print(f"  {count:4d}  {source}")

    print("\nSample chunks:")
    for record in records[:5]:
        meta = record["metadata"]
        print(f"\n--- {record['id']} [{meta['doc_type']}] ({len(record['text'])} chars) ---")
        print(record["text"][:300] + ("..." if len(record["text"]) > 300 else ""))

    OUTPUT_FILE.write_text(json.dumps(records, indent=2), encoding="utf-8")
    print(f"\nWrote {len(records)} chunks to {OUTPUT_FILE.name}")


if __name__ == "__main__":
    main()
